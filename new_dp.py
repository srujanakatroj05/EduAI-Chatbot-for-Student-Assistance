import os
import shutil
from docx import Document


def extract_text(doc):
    """Extracts text, headings, lists, footnotes, and hyperlinks from a DOCX file and formats them in Markdown."""
    md_text = []
    for para in doc.paragraphs:
        style = para.style.name
        text = para.text.strip()
        
        if not text:
            continue  # Skip empty lines
        
        if style.startswith("Heading"):
            level = int(style[-1]) if style[-1].isdigit() else 1
            md_text.append(f"{'#' * level} {text}")
        else:
            md_text.append(text)
    
    return "\n\n".join(md_text)


def extract_tables(doc):
    """Extracts tables from DOCX and converts them to Markdown format."""
    md_tables = []
    for table in doc.tables:
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append("| " + " | ".join(cells) + " |")
        
        if len(rows) > 1:
            header = rows[0]
            separator = "| " + " | ".join(["---"] * (len(rows[0].split("|")) - 2)) + " |"
            rows.insert(1, separator)
        
        md_tables.append("\n".join(rows))
    
    return "\n\n".join(md_tables)


def extract_images(doc, output_dir, doc_name):
    """Extracts images from DOCX and saves them in a single directory while adding Markdown references."""
    img_dir = os.path.join(output_dir, "images_output")  # Single folder for all images
    os.makedirs(img_dir, exist_ok=True)
    md_images = []
    img_count = 1
    
    for rel in doc.part.rels:
        if "image" in doc.part.rels[rel].target_ref:
            img_ext = doc.part.rels[rel].target_ref.split(".")[-1]
            img_filename = f"{doc_name}_image_{img_count}.{img_ext}"  # Unique filename
            img_path = os.path.join(img_dir, img_filename)
            
            with open(img_path, "wb") as img_file:
                img_file.write(doc.part.rels[rel].target_part.blob)
            
            md_images.append(f"![Image {img_count}](images_output/{img_filename})")
            img_count += 1
    
    return "\n\n".join(md_images)


def convert_docx_to_markdown(docx_path, output_md_path, output_dir):
    """Reads a DOCX file, extracts content, and saves it as Markdown."""
    doc_name = os.path.splitext(os.path.basename(docx_path))[0]
    doc = Document(docx_path)
    text_md = extract_text(doc)
    tables_md = extract_tables(doc)
    images_md = extract_images(doc, output_dir, doc_name)
    
    markdown_content = "\n\n".join([text_md, tables_md, images_md]).strip()
    
    with open(output_md_path, "w", encoding="utf-8") as f:
        f.write(markdown_content)
    
    print(f"Converted: {docx_path} → {output_md_path}")


def process_directory(input_dir, output_dir):
    """Processes all DOCX files in a directory and converts them to Markdown."""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)
    
    for filename in os.listdir(input_dir):
        file_path = os.path.join(input_dir, filename)
        file_name_without_ext = os.path.splitext(filename)[0]
        
        if filename.lower().endswith(".docx"):
            output_md_path = os.path.join(output_dir, f"{file_name_without_ext}.md")
            convert_docx_to_markdown(file_path, output_md_path, output_dir)


input_directory = "./docx_files_1/"     # "./docx_files/"   # Directory containing .docx files
output_directory = "./md_processed_2/"  # "./md_processed_1/"     # Where the .md files will be saved

process_directory(input_directory, output_directory)
